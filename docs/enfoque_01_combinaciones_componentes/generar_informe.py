"""
Genera el informe técnico del Enfoque 1 en formato Word (.docx).

Ejecutar desde la raíz del proyecto:
    python docs/enfoque_01_combinaciones_componentes/generar_informe.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT_PATH = Path(__file__).resolve().parent / "informe_tecnico_enfoque_01.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_bg(cell, hex_color: str):
    """Aplica color de fondo a una celda."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]],
              header_bg: str = "1565C0", col_widths: list[float] | None = None):
    """
    Crea una tabla con cabecera coloreada.
    col_widths: lista de anchos en cm para cada columna.
    """
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Cabecera
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Filas
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "F5F5F5" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(9)

    # Anchos de columna
    if col_widths:
        for i, row in enumerate(table.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Cm(col_widths[j])

    return table


def heading(doc: Document, text: str, level: int):
    p = doc.add_heading(text, level=level)
    if level == 1:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)  # azul oscuro
    elif level == 2:
        for run in p.runs:
            run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)
    return p


def body(doc: Document, text: str, bold_parts: list[str] | None = None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if bold_parts:
        # Divide el texto y aplica negrita a las partes indicadas
        remaining = text
        for part in bold_parts:
            idx = remaining.find(part)
            if idx >= 0:
                if idx > 0:
                    run = p.add_run(remaining[:idx])
                    run.font.size = Pt(10)
                bold_run = p.add_run(part)
                bold_run.bold = True
                bold_run.font.size = Pt(10)
                remaining = remaining[idx + len(part):]
        if remaining:
            run = p.add_run(remaining)
            run.font.size = Pt(10)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
    return p


def bullet(doc: Document, text: str):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10)
    return p


def code_block(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x21, 0x21, 0x21)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Fondo gris claro con sombreado de párrafo
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F3F3F3")
    pPr.append(shd)
    return p


# ---------------------------------------------------------------------------
# Construcción del documento
# ---------------------------------------------------------------------------

def build_document():
    doc = Document()

    # Márgenes
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

    # Estilo base de fuente
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # -----------------------------------------------------------------------
    # PORTADA
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("INFORME TÉCNICO")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(0x0D, 0x47, 0xA1)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Enfoque 1 — Combinaciones de Componentes Activos")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor(0x15, 0x65, 0xC0)

    doc.add_paragraph()
    meta_lines = [
        ("Asignatura:", "SCY1101 - Programación para la Ciencia de Datos"),
        ("Dataset:", "Medicine_Details.csv"),
        ("Institución:", "DuocUC"),
        ("Año:", "2026"),
    ]
    for label, value in meta_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run(f"{label} ")
        r1.bold = True
        r1.font.size = Pt(11)
        r2 = p.add_run(value)
        r2.font.size = Pt(11)

    doc.add_page_break()

    # -----------------------------------------------------------------------
    # 1. DESCRIPCIÓN DEL FOCO
    # -----------------------------------------------------------------------
    heading(doc, "1. Descripción del Foco de Análisis", 1)

    body(doc,
         "El Enfoque 1 estudia la estructura interna de las formulaciones farmacéuticas: "
         "qué principios activos componen cada medicamento, qué combinaciones son más frecuentes "
         "y cómo se relacionan entre sí en el mercado representado por el dataset.",
         bold_parts=["Enfoque 1"])

    body(doc,
         "La columna principal analizada es Composition, que contiene cadenas del tipo "
         "\"Amoxycillin (500mg) + Clavulanic Acid (125mg)\". El enfoque extrae los nombres "
         "de los componentes descartando las dosis y construye una red de co-ocurrencias "
         "para revelar los patrones de combinación.",
         bold_parts=["Composition"])

    heading(doc, "Preguntas de investigación", 2)

    questions = [
        ("¿Qué componentes individuales dominan el mercado?",
         "Metformin (664), Telmisartan (379), Glimepiride (379), Paracetamol (359)"),
        ("¿Qué pares de componentes son más frecuentes?",
         "Glimepiride + Metformin (325 medicamentos), dominando ampliamente"),
        ("¿El mercado está dominado por mono o multi-componentes?",
         "Sí: 59,8% monocomponente, 30,4% dúos. El 90,2% tiene 1–2 componentes"),
        ("¿Más componentes implica más efectos secundarios?",
         "No. Los complejos (5+) tienen la mediana de efectos más baja"),
        ("¿La valoración varía con el número de componentes?",
         "Los dúos (~41%) tienen la mejor valoración; decrece desde tríos en adelante"),
        ("¿Existen componentes hub?",
         "Sí. Metformin es el hub absoluto con 664 apariciones en el dataset"),
    ]

    add_table(doc,
              ["Pregunta", "Respuesta"],
              [[q, a] for q, a in questions],
              col_widths=[7.5, 9])

    # -----------------------------------------------------------------------
    # 2. ARCHIVOS DESARROLLADOS
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    heading(doc, "2. Archivos Desarrollados", 1)

    files = [
        ("cleaning.py", "Limpieza del dataset: duplicados, normalización, extracción de componentes, columnas derivadas"),
        ("transform.py", "Transformaciones: explode por componente, generación de pares, construcción de matriz de co-ocurrencia"),
        ("analysis.py", "Visualizaciones y exportación de gráficos (.png) y tablas (.csv)"),
        ("validation.py", "Validación de integridad: checksum MD5, esquema del DataFrame, comparación de shapes"),
    ]

    add_table(doc,
              ["Archivo", "Responsabilidad"],
              [[f"src/enfoque_01_combinaciones_componentes/{f}", r] for f, r in files],
              col_widths=[7.5, 9])

    # -----------------------------------------------------------------------
    # 3. PIPELINE IMPLEMENTADO
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    heading(doc, "3. Pipeline Implementado", 1)

    # 3.1 Limpieza
    heading(doc, "3.1 Limpieza (cleaning.py)", 2)
    body(doc,
         "Función principal: run_cleaning_pipeline(df). "
         "Ejecuta 5 pasos en orden con trazabilidad completa de cada decisión.",
         bold_parts=["run_cleaning_pipeline(df)"])

    pipeline_rows = [
        ["1", "eliminar_duplicados()", "Elimina 84 filas completamente idénticas"],
        ["2", "normalizar_composicion()", "Estandariza espacios y el separador + en Composition"],
        ["3", "extraer_componentes()", "Parsea nombres de principios activos, descarta dosis entre paréntesis, aplica Title Case"],
        ["4", "añadir_columnas_componentes()", "Genera components_list, n_components y size_category (pd.Categorical ordered=True)"],
        ["5", "flag_anomalies()", "Marca registros con n_components == 0 sin eliminarlos automáticamente"],
    ]
    add_table(doc, ["Paso", "Función", "Acción"], pipeline_rows, col_widths=[1.2, 5.5, 9.8])

    doc.add_paragraph()
    body(doc, "Columnas generadas por la limpieza:", bold_parts=["Columnas generadas por la limpieza:"])

    clean_cols = [
        ["components_list", "list[str]", "Lista de componentes sin dosis (ej: ['Amoxycillin', 'Clavulanic Acid'])"],
        ["n_components", "int", "Número de principios activos por medicamento"],
        ["size_category", "pd.Categorical", "Categoría ordinal: mono < duo < trio < cuádruple < complejo"],
        ["composition_anomaly", "bool", "True si no se extrajo ningún componente válido"],
    ]
    add_table(doc, ["Columna", "Tipo", "Descripción"], clean_cols, col_widths=[4, 3, 9.5])

    doc.add_paragraph()
    body(doc, "Impacto cuantificado de la limpieza:", bold_parts=["Impacto cuantificado de la limpieza:"])
    impact_rows = [
        ["Filas raw", "11.825"], ["Duplicados eliminados", "84 (−0,71%)"],
        ["Filas post-limpieza", "11.741"], ["Columnas añadidas", "4"],
        ["Componentes únicos detectados", "1.058"], ["Anomalías detectadas", "0"],
    ]
    add_table(doc, ["Métrica", "Valor"], impact_rows, col_widths=[7, 9.5])

    # 3.2 Transformaciones
    doc.add_paragraph()
    heading(doc, "3.2 Transformaciones (transform.py)", 2)
    body(doc,
         "Función principal: run_transform_pipeline(df_clean). "
         "Devuelve la tupla (df_exploded, df_pairs, cooc_matrix).",
         bold_parts=["run_transform_pipeline(df_clean)", "(df_exploded, df_pairs, cooc_matrix)"])

    transform_rows = [
        ["explotar_componentes()", "df_exploded",
         "Explode de components_list → una fila por componente (17.957 filas). Permite contar frecuencias individuales."],
        ["explotar_pares()", "df_pairs",
         "Genera todos los pares C(N,2) por medicamento con N≥2 componentes usando itertools.combinations. 8.227 pares totales."],
        ["construir_matriz_coocurrencia()", "cooc_matrix",
         "pd.crosstab() → matriz 20×20 simétrica. cooc[i][j] = número de medicamentos que contienen ambos componentes i y j."],
    ]
    add_table(doc, ["Función", "Salida", "Descripción"],
              transform_rows, col_widths=[4.5, 2.5, 9.5])

    doc.add_paragraph()
    body(doc,
         "Decisión técnica: los pares se ordenan alfabéticamente antes de generarse "
         "(sorted(components)), lo que garantiza que (A, B) y (B, A) siempre se "
         "representen de la misma forma, evitando duplicados espejo en el dataset de pares.",
         bold_parts=["sorted(components)"])

    body(doc, "Los tres DataFrames se exportan a data/processed/:",
         bold_parts=["data/processed/"])
    for f in ["medicine_exploded.csv", "medicine_pairs.csv", "cooc_matrix.csv"]:
        bullet(doc, f)

    # 3.3 Validación
    doc.add_paragraph()
    heading(doc, "3.3 Validación de Integridad (validation.py)", 2)
    body(doc,
         "Función principal: run_validation_pipeline(df_raw, df_clean). "
         "Exporta el reporte a outputs/reports/reporte_integridad.json.",
         bold_parts=["run_validation_pipeline(df_raw, df_clean)",
                     "outputs/reports/reporte_integridad.json"])

    val_rows = [
        ["1", "verificar_checksum()", "MD5 del CSV raw", "93656e73e4e8fbf7dd8da9ecfab7ce07 ✅"],
        ["2", "validar_esquema()", "Columnas y tipos del DataFrame raw", "9 columnas con tipos correctos ✅"],
        ["3", "comparar_shapes()", "Dimensiones antes/después", "84 filas eliminadas (0,71%) ✅"],
        ["4", "validar_columnas_clean()", "Columnas derivadas en df_clean", "Las 4 columnas presentes ✅"],
    ]
    add_table(doc, ["#", "Función", "Verificación", "Resultado"], val_rows,
              col_widths=[0.8, 4.5, 5, 6.2])

    # 3.4 Análisis
    doc.add_paragraph()
    heading(doc, "3.4 Análisis y Visualizaciones (analysis.py)", 2)
    body(doc,
         "Función principal: run_analysis_pipeline(df_clean, df_exploded, df_pairs, cooc_matrix). "
         "Genera 7 gráficos en outputs/figures/ y 2 tablas CSV en outputs/tables/.",
         bold_parts=["run_analysis_pipeline(df_clean, df_exploded, df_pairs, cooc_matrix)",
                     "outputs/figures/", "outputs/tables/"])

    vis_rows = [
        ["1", "plot_histograma_componentes()", "histograma_n_componentes.png",
         "¿Cómo se distribuyen los medicamentos por número de componentes?"],
        ["2", "plot_top_componentes()", "top_componentes_individuales.png",
         "¿Qué principios activos son más frecuentes en el dataset?"],
        ["3", "plot_top_pares()", "top_pares_componentes.png",
         "¿Qué combinaciones de dos componentes dominan?"],
        ["4", "plot_heatmap_coocurrencia()", "heatmap_coocurrencia.png",
         "¿Qué pares co-ocurren más entre los 20 componentes top?"],
        ["5", "plot_efectos_secundarios_por_tamanio()", "boxplot_efectos_secundarios.png",
         "¿Los medicamentos con más componentes tienen más efectos secundarios?"],
        ["6", "plot_network_graph()", "network_graph_coocurrencia.png",
         "¿Cómo se conectan los componentes en una red de co-ocurrencias?"],
        ["7", "plot_scatter_valoracion()", "scatter_valoracion_componentes.png",
         "¿La valoración de usuarios varía con el número de componentes?"],
    ]
    add_table(doc, ["#", "Función", "Archivo generado", "Pregunta respondida"],
              vis_rows, col_widths=[0.8, 5, 5, 5.7])

    # -----------------------------------------------------------------------
    # 4. RESULTADOS PRINCIPALES
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    heading(doc, "4. Resultados Principales", 1)

    # 4.1 Distribución
    heading(doc, "4.1 Distribución por Tamaño de Combinación", 2)
    body(doc,
         "El dataset presenta una distribución de cola larga fuertemente concentrada en "
         "monocomponentes y dúos. El 90,2% del mercado representado tiene 1 o 2 principios activos.",
         bold_parts=["90,2%"])

    dist_rows = [
        ["1 (mono)", "7.019", "59,8%"],
        ["2 (duo)", "3.569", "30,4%"],
        ["3 (trio)", "929", "7,9%"],
        ["4 (cuádruple)", "147", "1,3%"],
        ["5+ (complejo)", "77", "0,7%"],
        ["Total", "11.741", "100%"],
    ]
    add_table(doc, ["N° Componentes", "Medicamentos", "Porcentaje"],
              dist_rows, col_widths=[5, 5, 6.5])

    # 4.2 Top componentes
    doc.add_paragraph()
    heading(doc, "4.2 Componentes Hub — Top 10", 2)
    body(doc,
         "Metformin es el componente hub dominante con 664 apariciones, "
         "un 75% más que el segundo lugar. Su dominio refleja la alta prevalencia "
         "de diabetes tipo 2 en India, país de origen del dataset.",
         bold_parts=["Metformin", "664 apariciones"])

    comp_rows = [
        ["1", "Metformin", "664", "Antidiabético"],
        ["2", "Telmisartan", "379", "Antihipertensivo"],
        ["3", "Glimepiride", "379", "Antidiabético"],
        ["4", "Paracetamol", "359", "Analgésico/Antipirético"],
        ["5", "Amlodipine", "314", "Antihipertensivo"],
        ["6", "Montelukast", "224", "Antiasmático"],
        ["7", "Pregabalin", "214", "Antiepiléptico/Analgésico"],
        ["8", "Methylcobalamin", "214", "Vitamina B12"],
        ["9", "Metoprolol Succinate", "209", "Betabloqueante"],
        ["10", "Rosuvastatin", "209", "Estatina"],
    ]
    add_table(doc, ["Pos.", "Componente", "Medicamentos", "Área Terapéutica"],
              comp_rows, col_widths=[1.2, 5.5, 3.5, 6.3])

    # 4.3 Top pares
    doc.add_paragraph()
    heading(doc, "4.3 Pares de Componentes Más Frecuentes — Top 10", 2)
    body(doc,
         "Glimepiride + Metformin domina con 325 medicamentos, casi 2,6 veces el segundo par. "
         "Esta combinación une dos mecanismos complementarios para la diabetes tipo 2: "
         "reducción hepática de glucosa (Metformin) + estimulación pancreática de insulina (Glimepiride).",
         bold_parts=["Glimepiride + Metformin", "325 medicamentos"])

    pairs_rows = [
        ["1", "Glimepiride + Metformin", "325", "Antidiabético"],
        ["2", "Levocetirizine + Montelukast", "124", "Alérgico/Asma"],
        ["3", "Metformin + Voglibose", "124", "Antidiabético"],
        ["4", "Methylcobalamin + Pregabalin", "122", "Neuropatía"],
        ["5", "Amoxycillin + Clavulanic Acid", "97", "Antibiótico"],
        ["6", "Ambroxol + Guaifenesin", "91", "Respiratorio"],
        ["7", "Amlodipine + Telmisartan", "90", "Antihipertensivo"],
        ["8", "Aceclofenac + Paracetamol", "88", "Analgésico"],
        ["9", "Metformin + Pioglitazone", "87", "Antidiabético"],
        ["10", "Glimepiride + Voglibose", "86", "Antidiabético"],
    ]
    add_table(doc, ["Pos.", "Par", "Medicamentos", "Área Terapéutica"],
              pairs_rows, col_widths=[1.2, 7.5, 3.5, 4.3])

    # 4.4 Co-ocurrencia
    doc.add_paragraph()
    heading(doc, "4.4 Clusters de Co-ocurrencia", 2)
    body(doc,
         "La matriz de co-ocurrencia (20×20) y el grafo de red revelan tres clusters "
         "terapéuticos bien separados. Los componentes prácticamente no cruzan clusters, "
         "lo que valida la coherencia clínica del dataset.")

    cluster_rows = [
        ["Antidiabético", "Metformin, Glimepiride, Voglibose",
         "325 (Metformin–Glimepiride)", "Conexión más fuerte del dataset"],
        ["Antihipertensivo", "Telmisartan, Amlodipine, Hydrochlorothiazide, Olmesartan",
         "90 (Amlodipine–Telmisartan)", "Telmisartan actúa como hub del cluster"],
        ["Analgésico/Resp.", "Paracetamol, Aceclofenac, Phenylephrine, Chlorpheniramine",
         "88 (Paracetamol–Aceclofenac)", "Paracetamol es el nodo de mayor grado"],
    ]
    add_table(doc, ["Cluster", "Componentes", "Co-ocurrencia máx.", "Observación"],
              cluster_rows, col_widths=[3, 5, 4, 4.5])

    # 4.5 Efectos secundarios
    doc.add_paragraph()
    heading(doc, "4.5 Efectos Secundarios según Tamaño de Combinación", 2)
    body(doc,
         "Hallazgo contraintuitivo: los medicamentos complejos (5+ componentes) tienen la "
         "mediana más BAJA de efectos secundarios listados (~4), mientras que los tríos "
         "tienen la más alta (~7). Esto refuta la hipótesis inicial de que más componentes "
         "implica más efectos secundarios.",
         bold_parts=["más BAJA", "contraintuitivo"])

    efec_rows = [
        ["mono (1)", "6", "5 – 9"],
        ["duo (2)", "6", "5 – 9"],
        ["trio (3)", "7", "5 – 11"],
        ["cuádruple (4)", "5", "4 – 10"],
        ["complejo (5+)", "4", "3 – 5"],
    ]
    add_table(doc, ["Categoría", "Mediana efectos", "Rango IQR"],
              efec_rows, col_widths=[5, 5, 6.5])

    # 4.6 Valoración
    doc.add_paragraph()
    heading(doc, "4.6 Valoración de Usuarios por Número de Componentes", 2)
    body(doc,
         "Los dúos obtienen la mejor valoración promedio (~41%), superando incluso a los "
         "monocomponentes (~38%). Desde 3 componentes en adelante la valoración decrece de "
         "forma consistente. Los grupos de 7, 8 y 9 componentes tienen muestras "
         "insuficientes para conclusiones robustas.",
         bold_parts=["dúos", "~41%"])

    val_rows2 = [
        ["1 (mono)", "37,3%"],
        ["2 (duo)", "41,0% ← máximo"],
        ["3 (trio)", "38,9%"],
        ["4 (cuádruple)", "35,8%"],
        ["5 (complejo)", "35,6%"],
        ["6", "25,4%"],
    ]
    add_table(doc, ["N° Componentes", "Excellent Review % promedio"],
              val_rows2, col_widths=[7, 9.5])

    # -----------------------------------------------------------------------
    # 5. CONCLUSIONES
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    heading(doc, "5. Conclusiones", 1)

    conclusions = [
        ("Concentración del mercado:",
         "El 90,2% del dataset corresponde a formulaciones de 1 o 2 componentes. "
         "Los medicamentos de 5+ componentes son estadísticamente marginales (0,7%) "
         "y no deben dominar las conclusiones generales."),
        ("Dominio terapéutico:",
         "Los hubs del dataset son Metformin, Glimepiride, Telmisartan y Amlodipine, "
         "reflejando que el mercado farmacéutico representado está centrado en "
         "diabetes tipo 2, hipertensión arterial y dolor crónico."),
        ("Estructura de red:",
         "La red de co-ocurrencias forma tres clusters terapéuticos bien definidos "
         "y separados. No hay cruce de componentes entre áreas terapéuticas, "
         "lo cual valida la calidad y coherencia clínica del dataset."),
        ("Efectos secundarios:",
         "La complejidad de la formulación no aumenta linealmente los efectos secundarios. "
         "Las formulaciones complejas están más especializadas y tienen perfiles "
         "de seguridad más controlados."),
        ("Valoración de usuario:",
         "Los dúos representan el punto óptimo: mayor volumen, co-ocurrencias claras "
         "y mejor valoración promedio. La complejidad adicional (3+ componentes) "
         "no mejora la percepción del usuario."),
        ("Reproducibilidad:",
         "El pipeline es completamente trazable: checksum MD5 del archivo fuente, "
         "reporte JSON de validación, y exportación automática de todos los "
         "DataFrames intermedios a data/processed/."),
    ]

    for title, text in conclusions:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{title} ")
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = p.add_run(text)
        r2.font.size = Pt(10)

    # -----------------------------------------------------------------------
    # 6. EJECUCIÓN DEL PIPELINE
    # -----------------------------------------------------------------------
    doc.add_paragraph()
    heading(doc, "6. Instrucciones de Ejecución", 1)

    body(doc, "Para ejecutar el pipeline completo desde Python:")
    code_block(doc, "from src.load_data import load_medicine_data")
    code_block(doc, "from src.enfoque_01_combinaciones_componentes.cleaning import run_cleaning_pipeline")
    code_block(doc, "from src.enfoque_01_combinaciones_componentes.transform import run_transform_pipeline")
    code_block(doc, "from src.enfoque_01_combinaciones_componentes.analysis import run_analysis_pipeline")
    code_block(doc, "from src.enfoque_01_combinaciones_componentes.validation import run_validation_pipeline")
    doc.add_paragraph()
    code_block(doc, "df_raw = load_medicine_data()")
    code_block(doc, "df_clean = run_cleaning_pipeline(df_raw)")
    code_block(doc, "df_exploded, df_pairs, cooc_matrix = run_transform_pipeline(df_clean)")
    code_block(doc, "run_analysis_pipeline(df_clean, df_exploded, df_pairs, cooc_matrix)")
    code_block(doc, "run_validation_pipeline(df_raw, df_clean)")

    doc.add_paragraph()
    body(doc, "Alternativamente, ejecutar los notebooks en orden desde:")
    bullet(doc, "notebooks/enfoque_01_combinaciones_componentes/01_eda_combinaciones_componentes.ipynb")
    bullet(doc, "notebooks/enfoque_01_combinaciones_componentes/02_limpieza_transformacion_combinaciones_componentes.ipynb")
    bullet(doc, "notebooks/enfoque_01_combinaciones_componentes/03_analisis_combinaciones_componentes.ipynb")

    # -----------------------------------------------------------------------
    # Guardar
    # -----------------------------------------------------------------------
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"[OK] Informe guardado en: {OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
